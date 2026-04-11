"""
Generate a .docx letter.

Primary path  — template-based:
  Uses a .docx file uploaded by the admin as a Jinja2 template (docxtpl).
  The template must contain the following placeholders:
    {{number}}         — outgoing letter number
    {{date}}           — letter date (DD.MM.YYYY)
    {{recipient}}      — recipient name (bold/right-aligned from template formatting)
    {{subject}}        — letter subject
    {{body}}           — letter body (inserted as a sub-document)
    {{signer_name}}    — signer full name
    {{executor_name}}  — executor full name
    {{executor_phone}} — executor phone (may be empty)

Fallback path — programmatic (legacy):
  Builds the document from scratch when no template is uploaded.
  Layout: header (logo + org details) | footer (banner) | body sections.
"""
import asyncio
import logging
import os
import tempfile
from datetime import date as date_type
from docx import Document
from docx.shared import Cm, Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .html_to_docx import html_to_docx

log = logging.getLogger(__name__)

# Reference .docx for Pandoc — defines fonts, spacing, no-border table styles
_REFERENCE_DOCX = os.path.join(os.path.dirname(__file__), '..', 'reference.docx')

UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "/app/uploads")


# ── Shared helpers ────────────────────────────────────────────────────────────

def _format_date(d) -> str:
    if isinstance(d, date_type):
        return d.strftime("%d.%m.%Y")
    return str(d)


def _emu_to_twips(emu: int) -> int:
    """Convert EMU to Word twips (dxa). 1 twip = 635 EMU."""
    return int(emu / 635)


# ── Template-based generation ─────────────────────────────────────────────────

async def _build_body_subdoc(html: str, tpl, content_width: int):
    """
    Convert HTML body to a docxtpl Subdoc.
    Primary: Pandoc + reference.docx  →  faithful WYSIWYG output.
    Fallback: custom html_to_docx     →  used if Pandoc is unavailable.
    Returns (subdoc_object, tmp_path_to_delete_after_render).
    """
    if not html:
        return "", None

    tmp_path: str | None = None
    try:
        tmp_path = await _pandoc_convert(html)
        _remove_table_borders(tmp_path)
        subdoc = tpl.new_subdoc(tmp_path)
        return subdoc, tmp_path
    except Exception as exc:
        log.warning("Pandoc unavailable (%s), falling back to html_to_docx", exc)
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # ── Fallback ─────────────────────────────────────────────────────────────
    body_doc = Document()
    for sn in ("Normal", "List Bullet", "List Number",
               "List Bullet 2", "List Number 2"):
        try:
            s = body_doc.styles[sn]
            s.font.name = "Roboto"
            s.font.size = Pt(11)
        except KeyError:
            pass
    normal = body_doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(12)
    normal.paragraph_format.space_after  = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    html_to_docx(html, body_doc, content_width,
                 font_name="Roboto", font_size_pt=11.0,
                 space_before_pt=12.0, space_after_pt=12.0)

    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    body_doc.save(tmp_path)
    return tpl.new_subdoc(tmp_path), tmp_path


async def _pandoc_convert(html: str) -> str:
    """Run Pandoc HTML→DOCX. Returns path to the output temp file."""
    tmp_html_fd, tmp_html = tempfile.mkstemp(suffix=".html")
    tmp_docx_fd, tmp_docx = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_html_fd)
    os.close(tmp_docx_fd)
    try:
        with open(tmp_html, "w", encoding="utf-8") as f:
            f.write(f'<html><head><meta charset="utf-8"></head>'
                    f'<body>{html}</body></html>')

        cmd = ["pandoc", "-f", "html", "-t", "docx", "-o", tmp_docx, tmp_html]
        if os.path.exists(_REFERENCE_DOCX):
            cmd += ["--reference-doc", _REFERENCE_DOCX]

        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        _, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)
        if proc.returncode != 0:
            raise RuntimeError(stderr.decode(errors="replace"))

        return tmp_docx
    finally:
        if os.path.exists(tmp_html):
            os.unlink(tmp_html)


def _remove_table_borders(docx_path: str) -> None:
    """Strip all borders from every table in a .docx file (in-place)."""
    doc = Document(docx_path)
    for table in doc.tables:
        tblPr = table._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tblPr)
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        bdr = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            t = OxmlElement(f"w:{edge}")
            t.set(qn("w:val"), "none")
            bdr.append(t)
        tblPr.append(bdr)
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    continue
                for old in tcPr.findall(qn("w:tcBorders")):
                    tcPr.remove(old)
    doc.save(docx_path)


async def _generate_from_template(letter, org, template_path: str) -> str:
    from docxtpl import DocxTemplate

    tpl = DocxTemplate(template_path)

    sec = Document(template_path).sections[0]
    content_width = _emu_to_twips(sec.page_width - sec.left_margin - sec.right_margin)

    body_sd, tmp_path = await _build_body_subdoc(letter.body or "", tpl, content_width)

    sender = getattr(letter, "sender_type", "ooo")

    if sender == "ip":
        signer_name    = (org.ip_signer_name or "")    if org else ""
        signer_role    = (org.ip_signer_role or "")    if org else ""
        org_name       = (org.ip_full_name or "")      if org else ""
        org_short_name = (org.ip_full_name or "")      if org else ""
        org_inn        = (f"ИНН {org.ip_inn}," if org and org.ip_inn else "")
        org_ogrn       = (f"ОГРНИП {org.ip_ogrnip}" if org and org.ip_ogrnip else "")
        org_account    = (f"Р/с {org.ip_account}" if org and org.ip_account else "")
        org_bank       = (org.ip_bank_name or "")      if org else ""
        org_corr       = (f"Корр. счёт {org.ip_corr_account}" if org and org.ip_corr_account else "")
        org_bik        = (f"БИК {org.ip_bik}" if org and org.ip_bik else "")
        org_address    = (f"Адрес: {org.ip_legal_address}" if org and org.ip_legal_address else "")
        org_phone      = (org.ip_phone or "")          if org else ""
    else:
        signer_name    = (org.signer_name or "")       if org else ""
        signer_role    = (org.signer_role or "")       if org else ""
        org_name       = (org.name or "")              if org else ""
        org_short_name = (org.short_name or "")        if org else ""
        org_inn        = (f"ИНН {org.inn}," if org and org.inn else "")
        org_ogrn       = (f"ОГРН {org.ogrn}" if org and org.ogrn else "")
        org_account    = (f"Р/с {org.account}" if org and org.account else "")
        org_bank       = (org.bank_name or "")         if org else ""
        org_corr       = (f"Корр. счёт {org.corr_account}" if org and org.corr_account else "")
        org_bik        = (f"БИК {org.bik}" if org and org.bik else "")
        org_address    = (f"Адрес: {org.legal_address}" if org and org.legal_address else "")
        org_phone      = (org.phone or "")             if org else ""

    context = {
        "number":         str(letter.number),
        "date":           _format_date(letter.letter_date),
        "recipient":      letter.recipient.name if letter.recipient else "",
        "subject":        letter.subject or "",
        "body":           body_sd,
        "signer_name":    signer_name,
        "signer_role":    signer_role,
        "org_name":       org_name,
        "org_short_name": org_short_name,
        "org_inn":        org_inn,
        "org_ogrn":       org_ogrn,
        "org_account":    org_account,
        "org_bank":       org_bank,
        "org_corr":       org_corr,
        "org_bik":        org_bik,
        "org_address":    org_address,
        "org_phone":      org_phone,
        "executor_name":  letter.creator.full_name if letter.creator else "",
        "executor_phone": (letter.creator.phone or "") if letter.creator else "",
    }

    tpl.render(context, autoescape=False)

    output_dir = os.path.join(UPLOAD_DIR, "letters", str(letter.id))
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"letter_{letter.number}.docx")
    tpl.save(output_path)

    # Clean up temp body file after saving
    if tmp_path and os.path.exists(tmp_path):
        os.unlink(tmp_path)

    return output_path


# ── Programmatic (fallback) generation ───────────────────────────────────────

def _clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "auto")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _clear_paragraph_borders(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right", "between", "bar"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        pBdr.append(tag)
    pPr.append(pBdr)


def _set_cell_valign_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def _build_header(section, org, sender_type: str):
    header = section.header
    header.is_linked_to_previous = False

    page_width   = section.page_width
    right_margin = section.right_margin
    logo_col     = Cm(5.5)
    # Table starts at physical page left (via tblInd offset) and must end
    # at the right margin boundary, not the physical page edge.
    table_width  = page_width - right_margin
    details_col  = table_width - logo_col

    htable = header.add_table(1, 2, table_width)
    tblPr = htable._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        htable._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tblBorders.append(tag)
    tblPr.append(tblBorders)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(-_emu_to_twips(section.left_margin)))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    left_cell  = htable.cell(0, 0)
    right_cell = htable.cell(0, 1)
    _clear_cell_borders(left_cell)
    _clear_cell_borders(right_cell)
    left_cell.width  = logo_col
    right_cell.width = details_col

    _set_cell_valign_center(left_cell)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _clear_paragraph_borders(left_para)
    if org and org.logo_path and os.path.exists(org.logo_path):
        run = left_para.add_run()
        run.add_picture(org.logo_path, width=Cm(5.87), height=Cm(3.13))

    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _clear_paragraph_borders(right_para)

    def _line(text: str, bold: bool = False):
        if not text:
            return
        run = right_para.add_run(("\n" if right_para.runs else "") + text)
        run.font.name  = "Roboto"
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x15, 0x15, 0x15)
        run.bold = bold

    if org:
        if sender_type == "ip":
            _line(org.ip_full_name)
            if org.ip_inn:        _line(f"ИНН {org.ip_inn}")
            if org.ip_ogrnip:     _line(f"ОГРНИП {org.ip_ogrnip}")
            if org.ip_account:    _line(f"Р/с {org.ip_account}")
            if org.ip_bank_name:  _line(org.ip_bank_name)
            if org.ip_corr_account: _line(f"К/с {org.ip_corr_account}")
            if org.ip_bik:        _line(f"БИК {org.ip_bik}")
            if org.ip_legal_address: _line(org.ip_legal_address)
            if org.ip_phone:      _line(f"Тел.: {org.ip_phone}")
        else:
            _line(org.name)
            if org.inn:           _line(f"ИНН {org.inn}")
            if org.ogrn:          _line(f"ОГРН {org.ogrn}")
            if org.account:       _line(f"Р/с {org.account}")
            if org.bank_name:     _line(org.bank_name)
            if org.corr_account:  _line(f"К/с {org.corr_account}")
            if org.bik:           _line(f"БИК {org.bik}")
            if org.legal_address: _line(org.legal_address)
            if org.phone:         _line(f"Тел.: {org.phone}")

    default_para = header.paragraphs[0]
    _clear_paragraph_borders(default_para)
    header._element.append(default_para._p)


def _build_footer(section, org):
    """
    Fill footer with an ANCHORED banner image (like the reference document).
    wp:anchor + wrapNone keeps the paragraph height minimal; image floats at
    the page bottom.
    """
    if not (org and org.footer_banner_path and os.path.exists(org.footer_banner_path)):
        return

    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pPr = para._p.get_or_add_pPr()
    sp_elem = OxmlElement("w:spacing")
    sp_elem.set(qn("w:before"), "0")
    sp_elem.set(qn("w:after"),  "0")
    pPr.append(sp_elem)

    run = para.add_run()
    run.add_picture(org.footer_banner_path, width=Cm(20.97), height=Cm(3.79))

    r_elem  = run._r
    drawing = r_elem.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return

    extent  = inline.find(qn("wp:extent"))
    effect  = inline.find(qn("wp:effectExtent"))
    docPr   = inline.find(qn("wp:docPr"))
    cNvGFP  = inline.find(qn("wp:cNvGraphicFramePr"))
    graphic = inline.find(qn("a:graphic"))

    anchor = OxmlElement("wp:anchor")
    anchor.set("distT", "0"); anchor.set("distB", "0")
    anchor.set("distL", "0"); anchor.set("distR", "0")
    anchor.set("simplePos",    "0")
    anchor.set("relativeHeight", "251653120")
    anchor.set("behindDoc",    "0")
    anchor.set("locked",       "0")
    anchor.set("layoutInCell", "1")
    anchor.set("allowOverlap", "0")

    sp_pos = OxmlElement("wp:simplePos")
    sp_pos.set("x", "0"); sp_pos.set("y", "0")
    anchor.append(sp_pos)

    posH = OxmlElement("wp:positionH")
    posH.set("relativeFrom", "column")
    oh = OxmlElement("wp:posOffset")
    oh.text = str(-int(section.left_margin))
    posH.append(oh); anchor.append(posH)

    posV = OxmlElement("wp:positionV")
    posV.set("relativeFrom", "paragraph")
    ov = OxmlElement("wp:posOffset")
    ov.text = "0"
    posV.append(ov); anchor.append(posV)

    for child in (extent, effect):
        if child is not None:
            anchor.append(child)
    anchor.append(OxmlElement("wp:wrapNone"))
    for child in (docPr, cNvGFP, graphic):
        if child is not None:
            anchor.append(child)

    drawing.remove(inline)
    drawing.append(anchor)


async def _generate_programmatic(letter, org) -> str:
    output_dir = os.path.join(UPLOAD_DIR, "letters", str(letter.id))
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"letter_{letter.number}.docx")

    doc = Document()
    section = doc.sections[0]

    section.top_margin    = Twips(1360)
    section.bottom_margin = Twips(0)
    section.left_margin   = Twips(1276)
    section.right_margin  = Twips(707)
    section.header_distance = Twips(720)
    section.footer_distance = Twips(291)

    style = doc.styles["Normal"]
    style.font.name = "Roboto"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    sender = getattr(letter, "sender_type", "ooo")

    _build_header(section, org, sender)
    _build_footer(section, org)

    # Meta line
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.add_run(f"Исх.№{letter.number} от {_format_date(letter.letter_date)} г.")

    # Recipient
    if letter.recipient:
        rec_para = doc.add_paragraph()
        rec_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = rec_para.add_run(letter.recipient.name)
        run.bold = True
        run.font.name = "Roboto"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Subject
    if letter.subject:
        subj_para = doc.add_paragraph()
        subj_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = subj_para.add_run(letter.subject)
        run.italic = True
        run.font.name = "Roboto"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()

    # Body
    content_width = _emu_to_twips(section.page_width - section.left_margin - section.right_margin)
    if letter.body:
        html_to_docx(letter.body, doc, content_width)

    doc.add_paragraph()

    # Signature table — span the full content width
    content_w = section.page_width - section.left_margin - section.right_margin
    sig_col1 = int(content_w * 0.44)   # "С уважением, роль организация"
    sig_col2 = int(content_w * 0.22)   # подпись (картинка)
    sig_col3 = content_w - sig_col1 - sig_col2  # ФИО подписанта

    sig_table = doc.add_table(rows=1, cols=3)
    for c in sig_table.rows[0].cells:
        _clear_cell_borders(c)
    sig_table.rows[0].cells[0].width = sig_col1
    sig_table.rows[0].cells[1].width = sig_col2
    sig_table.rows[0].cells[2].width = sig_col3

    sl = sig_table.cell(0, 0).paragraphs[0]
    sl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if sender == "ip":
        signer_role = org.ip_signer_role if org else ""
        short_name  = ""  # ФИО ИП уже в правой ячейке; роль содержит "ИП ФИО"
    else:
        signer_role = org.signer_role if org else ""
        short_name  = org.short_name  if org else ""
    left_text = "С уважением,"
    if signer_role: left_text += f"\n{signer_role}"
    if short_name:  left_text += f"\n{short_name}"
    run = sl.add_run(left_text)
    run.font.name = "Roboto"
    run.font.color.rgb = RGBColor(0, 0, 0)

    sm = sig_table.cell(0, 1).paragraphs[0]
    sm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if org and org.signature_path and os.path.exists(org.signature_path):
        run = sm.add_run()
        run.add_picture(org.signature_path, width=Cm(3))

    sr = sig_table.cell(0, 2).paragraphs[0]
    sr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signer_name = (org.ip_signer_name if sender == "ip" else org.signer_name) if org else ""
    if signer_name:
        run = sr.add_run(signer_name)
        run.font.name = "Roboto"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Executor
    doc.add_paragraph()
    exec_para = doc.add_paragraph()
    exec_text = "Исп.:"
    if letter.creator:
        exec_text += f"\n{letter.creator.full_name}"
        if letter.creator.phone:
            exec_text += f"\n{letter.creator.phone}"
    run = exec_para.add_run(exec_text)
    run.font.name = "Roboto"
    run.font.size = Pt(10)

    doc.save(output_path)
    return output_path


# ── Public entry point ────────────────────────────────────────────────────────

async def generate_letter_docx(letter, org) -> str:
    """
    Generate a letter .docx.
    Uses the org's uploaded template if available; falls back to programmatic build.
    """
    sender = getattr(letter, "sender_type", "ooo")
    template_path = None
    if org:
        template_path = org.template_ip_path if sender == "ip" else org.template_ooo_path

    if template_path and os.path.exists(template_path):
        return await _generate_from_template(letter, org, template_path)

    return await _generate_programmatic(letter, org)
