"""
Convert Tiptap HTML to python-docx elements.
Supports: p, strong, em, ul, ol, li, table/tr/td/th.
"""
import re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


def _add_run_with_formatting(paragraph, text: str, bold: bool = False, italic: bool = False,
                              font_name: str | None = None, font_size_pt: float | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_name:
        run.font.name = font_name
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    return run


def _set_para_spacing(paragraph, space_before_pt: float = 0.0, space_after_pt: float = 6.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after  = Pt(space_after_pt)


def _process_inline(element, paragraph, bold=False, italic=False, link=False,
                    font_name: str | None = None, font_size_pt: float | None = None):
    """Recursively process inline elements into a paragraph."""
    if isinstance(element, NavigableString):
        text = str(element)
        if text:
            run = _add_run_with_formatting(paragraph, text, bold=bold, italic=italic,
                                           font_name=font_name, font_size_pt=font_size_pt)
            if link:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.font.underline = True
    elif isinstance(element, Tag):
        is_bold   = bold   or element.name in ("strong", "b")
        is_italic = italic or element.name in ("em", "i")
        is_link   = link   or element.name == "a"
        for child in element.children:
            _process_inline(child, paragraph, bold=is_bold, italic=is_italic, link=is_link,
                            font_name=font_name, font_size_pt=font_size_pt)


def _clear_borders(obj_tbl_or_cell, is_cell: bool = False):
    """Remove all borders from a table or cell XML element."""
    tc_or_tbl = obj_tbl_or_cell._tc if is_cell else obj_tbl_or_cell._tbl
    pr_tag = "w:tcPr" if is_cell else "w:tblPr"
    border_tag = "w:tcBorders" if is_cell else "w:tblBorders"

    pPr = tc_or_tbl.find(qn(pr_tag))
    if pPr is None:
        pPr = OxmlElement(pr_tag)
        tc_or_tbl.insert(0, pPr)

    borders = OxmlElement(border_tag)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tag.set(qn("w:sz"), "0")
        tag.set(qn("w:color"), "auto")
        borders.append(tag)
    pPr.append(borders)


def _add_thin_borders(obj_tbl_or_cell, is_cell: bool = False):
    """Add thin single borders to a table or cell."""
    tc_or_tbl = obj_tbl_or_cell._tc if is_cell else obj_tbl_or_cell._tbl
    pr_tag    = "w:tcPr" if is_cell else "w:tblPr"
    border_tag = "w:tcBorders" if is_cell else "w:tblBorders"

    pPr = tc_or_tbl.find(qn(pr_tag))
    if pPr is None:
        pPr = OxmlElement(pr_tag)
        tc_or_tbl.insert(0, pPr)

    # Remove existing border element if present
    for old in pPr.findall(qn(border_tag)):
        pPr.remove(old)

    borders = OxmlElement(border_tag)
    edges = ("top", "left", "bottom", "right", "insideH", "insideV") if not is_cell \
            else ("top", "left", "bottom", "right")
    for edge in edges:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")       # 0.5pt
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "808080")
        borders.append(tag)
    pPr.append(borders)


def _extract_col_widths_px(tag: Tag, max_cols: int) -> list[float] | None:
    """
    Try to extract per-column pixel widths from Tiptap table HTML.
    Returns a list of floats (px) or None if no usable info found.
    """
    rows = tag.find_all("tr")
    if not rows:
        return None

    first_row = rows[0]
    cells = first_row.find_all(["td", "th"])

    widths: list[float | None] = []
    for cell in cells[:max_cols]:
        cw = cell.get("colwidth")
        if cw:
            try:
                widths.append(float(cw))
                continue
            except ValueError:
                pass
        style = cell.get("style", "")
        m = re.search(r"(?:min-)?width:\s*([\d.]+)px", style)
        if m:
            widths.append(float(m.group(1)))
            continue
        widths.append(None)

    if all(w is None for w in widths):
        colgroup = tag.find("colgroup")
        if colgroup:
            cols = colgroup.find_all("col")
            widths = []
            for col in cols[:max_cols]:
                style = col.get("style", "")
                attr  = col.get("width", "")
                m = re.search(r"(?:min-)?width:\s*([\d.]+)px", style)
                if m:
                    widths.append(float(m.group(1)))
                elif attr:
                    try:
                        widths.append(float(attr.replace("px", "").strip()))
                    except ValueError:
                        widths.append(None)
                else:
                    widths.append(None)

    while len(widths) < max_cols:
        widths.append(None)
    widths = widths[:max_cols]

    if all(w is None for w in widths):
        return None
    return widths


def _apply_table_layout(table, content_width: int, col_widths_px: list[float | None] | None, max_cols: int):
    """Set exact table width and column widths via tblGrid + per-cell tcW."""
    if col_widths_px and any(w is not None for w in col_widths_px):
        known = [w for w in col_widths_px if w is not None and w > 0]
        avg = sum(known) / len(known) if known else 1
        filled = [w if (w is not None and w > 0) else avg for w in col_widths_px]
        total_px = sum(filled)
        col_widths = [int(content_width * w / total_px) for w in filled]
        col_widths[-1] += content_width - sum(col_widths)
    else:
        base = content_width // max_cols
        col_widths = [base] * max_cols
        col_widths[-1] += content_width - sum(col_widths)

    tbl = table._tbl

    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(content_width))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.append(tblGrid)
    existing = tblGrid.findall(qn("w:gridCol"))
    for i, w in enumerate(col_widths):
        if i < len(existing):
            existing[i].set(qn("w:w"), str(w))
        else:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            tblGrid.append(gc)

    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j >= max_cols:
                break
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths[j]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _process_table(tag: Tag, document: Document, content_width: int | None = None,
                   font_name: str | None = None, font_size_pt: float | None = None,
                   space_before_pt: float = 0.0, space_after_pt: float = 3.0):
    rows = tag.find_all("tr")
    if not rows:
        return

    max_cols = max(len(r.find_all(["td", "th"])) for r in rows)
    if max_cols == 0:
        return

    table = document.add_table(rows=len(rows), cols=max_cols)
    _add_thin_borders(table)

    if content_width:
        col_widths_px = _extract_col_widths_px(tag, max_cols)
        _apply_table_layout(table, content_width, col_widths_px, max_cols)

    for i, row_tag in enumerate(rows):
        cells = row_tag.find_all(["td", "th"])
        for j, cell_tag in enumerate(cells):
            if j >= max_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            _add_thin_borders(cell, is_cell=True)
            para = cell.paragraphs[0]
            _set_para_spacing(para, space_before_pt, space_after_pt)
            is_header = cell_tag.name == "th"
            for child in cell_tag.children:
                _process_inline(child, para, bold=is_header,
                                font_name=font_name, font_size_pt=font_size_pt)


def html_to_docx(html: str, document: Document, content_width: int | None = None,
                 font_name: str | None = None, font_size_pt: float | None = None,
                 space_before_pt: float = 0.0, space_after_pt: float = 6.0):
    """Parse HTML and append content to the document."""
    if not html:
        return

    soup = BeautifulSoup(html, "html.parser")

    def _process_block(element):
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = document.add_paragraph()
                _set_para_spacing(p, space_before_pt, space_after_pt)
                _add_run_with_formatting(p, text, font_name=font_name, font_size_pt=font_size_pt)
            return

        name = element.name if isinstance(element, Tag) else None
        if name is None:
            return

        if name == "p":
            p = document.add_paragraph()
            _set_para_spacing(p, space_before_pt, space_after_pt)
            for child in element.children:
                _process_inline(child, p, font_name=font_name, font_size_pt=font_size_pt)

        elif name in ("ul", "ol"):
            style = "List Bullet" if name == "ul" else "List Number"
            for li in element.find_all("li", recursive=False):
                p = document.add_paragraph(style=style)
                _set_para_spacing(p, space_before_pt, space_after_pt)
                for child in li.children:
                    _process_inline(child, p, font_name=font_name, font_size_pt=font_size_pt)

        elif name == "table":
            _process_table(element, document, content_width,
                           font_name=font_name, font_size_pt=font_size_pt,
                           space_before_pt=0.0, space_after_pt=3.0)

        elif name in ("div", "section", "article"):
            for child in element.children:
                _process_block(child)

        elif name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            p = document.add_paragraph()
            _set_para_spacing(p, space_before_pt, space_after_pt)
            for child in element.children:
                _process_inline(child, p, bold=True, font_name=font_name, font_size_pt=font_size_pt)

        elif name == "br":
            p = document.add_paragraph()
            _set_para_spacing(p, 0, 0)

        else:
            p = document.add_paragraph()
            _set_para_spacing(p, space_before_pt, space_after_pt)
            for child in element.children:
                _process_inline(child, p, font_name=font_name, font_size_pt=font_size_pt)

    for child in soup.children:
        _process_block(child)
